import os
import logging
from typing import Optional, Dict, Any
from pathlib import Path

try:
    import PyPDF2
except ImportError:
    PyPDF2 = None

try:
    from docx import Document
except ImportError:
    Document = None

logger = logging.getLogger(__name__)

class FileProcessor:
    """Utility class for processing different file types"""
    
    def __init__(self):
        self.supported_types = {
            '.md': self._process_markdown,
            '.txt': self._process_text,
            '.pdf': self._process_pdf,
            '.docx': self._process_docx
        }
    
    async def process_file(self, filename: str, content: bytes) -> Dict[str, Any]:
        """Process file based on its extension"""
        try:
            file_ext = Path(filename).suffix.lower()
            
            if file_ext not in self.supported_types:
                raise ValueError(f"Unsupported file type: {file_ext}. Supported types: {list(self.supported_types.keys())}")
            
            processor = self.supported_types[file_ext]
            text_content = await processor(content)
            
            return {
                'text_content': text_content,
                'file_type': file_ext,
                'original_size': len(content),
                'processed_size': len(text_content)
            }
            
        except Exception as e:
            logger.error(f"Error processing file {filename}: {str(e)}")
            raise
    
    async def _process_markdown(self, content: bytes) -> str:
        """Process markdown files"""
        return content.decode('utf-8')
    
    async def _process_text(self, content: bytes) -> str:
        """Process text files"""
        return content.decode('utf-8')
    
    async def _process_pdf(self, content: bytes) -> str:
        """Process PDF files"""
        if PyPDF2 is None:
            raise ImportError("PyPDF2 is required for PDF processing. Install with: pip install PyPDF2")
        
        try:
            from io import BytesIO
            pdf_file = BytesIO(content)
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            
            text_content = ""
            for page_num, page in enumerate(pdf_reader.pages):
                try:
                    page_text = page.extract_text()
                    if page_text.strip():
                        text_content += f"\n\n--- Page {page_num + 1} ---\n\n"
                        text_content += page_text
                except Exception as e:
                    logger.warning(f"Error extracting text from page {page_num + 1}: {str(e)}")
                    continue
            
            if not text_content.strip():
                raise ValueError("No text content could be extracted from the PDF")
            
            return text_content.strip()
            
        except Exception as e:
            logger.error(f"Error processing PDF: {str(e)}")
            raise ValueError(f"Failed to process PDF file: {str(e)}")
    
    async def _process_docx(self, content: bytes) -> str:
        """Process DOCX files"""
        if Document is None:
            raise ImportError("python-docx is required for DOCX processing. Install with: pip install python-docx")
        
        try:
            from io import BytesIO
            docx_file = BytesIO(content)
            doc = Document(docx_file)
            
            text_content = ""
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content += paragraph.text + "\n\n"
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_content += " | ".join(row_text) + "\n"
            
            if not text_content.strip():
                raise ValueError("No text content could be extracted from the DOCX file")
            
            return text_content.strip()
            
        except Exception as e:
            logger.error(f"Error processing DOCX: {str(e)}")
            raise ValueError(f"Failed to process DOCX file: {str(e)}")
    
    def is_supported_file(self, filename: str) -> bool:
        """Check if file type is supported"""
        file_ext = Path(filename).suffix.lower()
        return file_ext in self.supported_types
    
    def get_supported_extensions(self) -> list:
        """Get list of supported file extensions"""
        return list(self.supported_types.keys())